"""Générateur de documents Word pour l'Agent RH.

Utilise python-docx pour prendre un template Word (issu de Google Drive)
et le remplir avec les données générées par le LLM (CV personnalisé,
lettre de motivation, etc.).
"""

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def load_template(template_path: Path) -> Document:
    """Charge un template Word (.docx) depuis le disque.

    Args:
        template_path: Chemin vers le fichier template .docx.

    Returns:
        Objet Document python-docx.

    Raises:
        FileNotFoundError: Si le template n'existe pas.
    """
    if not template_path.exists():
        raise FileNotFoundError(f"Template introuvable : {template_path}")
    return Document(str(template_path))


def _replace_in_paragraph(paragraph, placeholders: dict[str, str]) -> None:
    """Remplace les placeholders dans un paragraphe en préservant le style.

    Reconstruit le texte complet du paragraphe, effectue les remplacements,
    puis réattribue le texte au premier run et vide les autres.
    """
    full_text = "".join(run.text for run in paragraph.runs)
    if not re.search(r"\{\{.+?\}\}", full_text):
        return

    for key, value in placeholders.items():
        full_text = full_text.replace(f"{{{{{key}}}}}", value)

    # Réattribuer le texte modifié au premier run, vider les autres
    for i, run in enumerate(paragraph.runs):
        if i == 0:
            run.text = full_text
        else:
            run.text = ""


def fill_template(
    template_path: Path,
    placeholders: dict[str, str],
    output_path: Path,
) -> Path:
    """Remplit un template Word en remplaçant les placeholders.

    Les placeholders dans le document sont au format {{NOM_VARIABLE}}.
    Le remplacement s'effectue dans les paragraphes, les en-têtes,
    les pieds de page et les tableaux.

    Args:
        template_path: Chemin du template source.
        placeholders: Dictionnaire {placeholder: valeur de remplacement}.
        output_path: Chemin de sortie du document généré.

    Returns:
        Chemin du document généré.
    """
    doc = load_template(template_path)

    # Paragraphes du corps
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, placeholders)

    # Tableaux
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, placeholders)

    # En-têtes et pieds de page
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            _replace_in_paragraph(paragraph, placeholders)
        for paragraph in section.footer.paragraphs:
            _replace_in_paragraph(paragraph, placeholders)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def generate_cover_letter_docx(
    cover_letter_text: str,
    candidate_name: str,
    company_name: str,
    output_path: Path,
) -> Path:
    """Génère un document Word contenant la lettre de motivation.

    Crée un document formaté de zéro (sans template) avec en-tête,
    date, corps de lettre et signature.

    Args:
        cover_letter_text: Texte de la lettre de motivation.
        candidate_name: Nom complet du candidat.
        company_name: Nom de l'entreprise destinataire.
        output_path: Chemin de sortie du fichier .docx.

    Returns:
        Chemin du document généré.
    """
    doc = Document()

    # -- Style par défaut --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # -- En-tête : nom du candidat --
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(candidate_name)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # -- Date --
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(datetime.now().strftime("%d/%m/%Y"))
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # -- Destinataire --
    dest_para = doc.add_paragraph()
    dest_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    dest_run = dest_para.add_run(f"À l'attention du service recrutement\n{company_name}")
    dest_run.font.size = Pt(10)
    dest_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # -- Espacement --
    doc.add_paragraph()

    # -- Corps de la lettre --
    for paragraph_text in cover_letter_text.split("\n\n"):
        paragraph_text = paragraph_text.strip()
        if not paragraph_text:
            continue
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(paragraph_text)
        run.font.size = Pt(11)

    # -- Signature --
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sig_run = sig_para.add_run(f"Cordialement,\n{candidate_name}")
    sig_run.font.size = Pt(11)
    sig_run.bold = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
